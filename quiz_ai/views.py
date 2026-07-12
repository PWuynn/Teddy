from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Quiz, Question, Choice, QuizResult
from django.http import HttpResponseForbidden
from django.utils import timezone
from datetime import datetime
from django.contrib import messages
from django.db.models import Q
from django.db import transaction
from classroom.models import ClassroomMember
from django.template.loader import get_template
from django.core.files.base import ContentFile
import re
import unicodedata
import random
from io import BytesIO


_CHOICE_MARKER_RE = re.compile(r'(?<!\w)[\s\-*]*(?:\(?[A-Fa-f]\)?)[\s]*[\.\):\-/\u2013\u2014]\s*')
_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*(?:\(?[A-Fa-f]\)?)[\s]*(?:[\.\):\-/\u2013\u2014]|\s+)')
_NAMED_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*(?:dap an|phuong an|option|choice)\s+([a-f])\s*[\.\):\-/\u2013\u2014]\s+')
_NUMERIC_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*[1-9][0-9]*[\.)]\s*\S+$')
_NUMERIC_CHOICE_PREFIX_RE = re.compile(r'^[\s\-*]*[1-9][0-9]*[\.)]\s*')
_BULLET_CHOICE_LINE_RE = re.compile(r'^[\s]*(?:[-\u2022\u2013\u2014]|\*)\s+\S+$')
_BULLET_CHOICE_PREFIX_RE = re.compile(r'^[\s]*(?:[-\u2022\u2013\u2014]|\*)\s+')
_ANSWER_PREFIXES = (
    'dap an:', 'dap an dung:', 'answer:', 'answer key:', 'correct:',
)


def _fold_quiz_text(value):
    normalized = unicodedata.normalize('NFD', value.lower().strip())
    without_marks = ''.join(ch for ch in normalized if unicodedata.category(ch) != 'Mn')
    return without_marks.replace(chr(273), 'd')


def _decode_uploaded_text(uploaded_file):
    raw_content = uploaded_file.read()
    for encoding in ('utf-8-sig', 'utf-8', 'cp1258', 'cp1252'):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_content.decode('utf-8', errors='replace')


def _normalize_quiz_text(value):
    return value.strip().replace('<br />', '\n').replace('<br/>', '\n').replace('<br>', '\n')



def _decode_uploaded_bytes(raw_content):
    for encoding in ('utf-8-sig', 'utf-8', 'cp1258', 'cp1252'):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_content.decode('utf-8', errors='replace')



def _extract_uploaded_quiz_images(uploaded_file):
    """Return embedded Word images in document order for automatic quizzes."""
    filename = (getattr(uploaded_file, 'name', '') or '').lower()
    if not filename.endswith('.docx'):
        return []

    try:
        from docx import Document
        raw_content = uploaded_file.read()
        uploaded_file.seek(0)
        document = Document(BytesIO(raw_content))
    except Exception:
        return []

    images = []
    extension_by_type = {
        'image/jpeg': 'jpg', 'image/png': 'png', 'image/gif': 'gif',
        'image/webp': 'webp', 'image/bmp': 'bmp',
    }
    for index, shape in enumerate(document.inline_shapes, start=1):
        try:
            relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = document.part.related_parts[relationship_id]
            extension = extension_by_type.get(image_part.content_type, 'png')
            images.append(ContentFile(image_part.blob, name=f'question_{index}.{extension}'))
        except Exception:
            continue
    return images


def _extract_uploaded_quiz_text(uploaded_file):
    raw_content = uploaded_file.read()
    try:
        uploaded_file.seek(0)
    except (AttributeError, OSError):
        pass

    filename = (getattr(uploaded_file, 'name', '') or '').lower()
    extension = filename.rsplit('.', 1)[-1] if '.' in filename else ''

    if extension in {'txt', 'text', 'csv'}:
        return _decode_uploaded_bytes(raw_content), []

    if extension == 'docx':
        try:
            from docx import Document
        except ImportError:
            return '', ['Máy chủ chưa có thư viện đọc Word. Vui lòng cài python-docx.']
        try:
            document = Document(BytesIO(raw_content))
        except Exception:
            return '', ['Không thể đọc tệp Word. Vui lòng tải tệp .docx hợp lệ.']

        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 2 or not cells[0] or not cells[1]:
                    continue
                header = _fold_quiz_text(f'{cells[0]} {cells[1]}')
                if 'cau hoi' in header and ('dap an' in header or 'lua chon' in header):
                    continue
                blocks.append(f'{cells[0]}\n{cells[1]}')
        # Each Word paragraph is a consecutive quiz line. Separating every
        # paragraph with a blank line flushes the parser before its choices.
        return '\n'.join(blocks), []

    if extension == 'pdf':
        try:
            import pdfplumber
        except ImportError:
            return '', ['Máy chủ chưa có thư viện đọc PDF. Vui lòng cài pdfplumber.']
        try:
            with pdfplumber.open(BytesIO(raw_content)) as pdf:
                table_blocks = []
                plain_text = []
                for page in pdf.pages:
                    plain_text.append(page.extract_text() or '')
                    for table in page.extract_tables() or []:
                        for row in table:
                            cells = [(cell or '').strip() for cell in row]
                            if len(cells) < 2 or not cells[0] or not cells[1]:
                                continue
                            header = _fold_quiz_text(f'{cells[0]} {cells[1]}')
                            if 'cau hoi' in header and ('dap an' in header or 'lua chon' in header):
                                continue
                            table_blocks.append(f'{cells[0]}\n{cells[1]}')
        except Exception:
            return '', ['Không thể đọc tệp PDF. Vui lòng tải PDF có thể chọn/sao chép văn bản.']
        return ('\n\n'.join(table_blocks) if table_blocks else '\n'.join(plain_text)), []

    return '', ['Chỉ hỗ trợ tệp .docx, .pdf, .txt, .text hoặc .csv.']


def _split_manual_question_blocks(text):
    normalized = text.replace('\r\n', '\n').replace('\r', '\n')
    blocks = []
    current = []

    for line in normalized.split('\n'):
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
            continue
        if stripped.startswith("'"):
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(stripped)

    if current:
        blocks.append(current)

    return blocks


def _extract_correct_letters(lines):
    correct_letters = set()
    clean_lines = []

    for line_number, line in lines:
        lowered = _fold_quiz_text(line)
        if lowered.startswith(_ANSWER_PREFIXES):
            answer_value = line.split(':', 1)[1].strip() if ':' in line else ''
            compact_key = re.sub(r'[\s,;/&]+', '', _fold_quiz_text(answer_value))
            compact_key = compact_key.replace('va', '').replace('and', '')
            if compact_key and all(letter in 'abcdef' for letter in compact_key):
                correct_letters.update(compact_key)
                continue
            # A line such as "Đáp án: Hà Nội" is itself the only choice,
            # not an answer-key label. Keep its value as the choice text.
            clean_lines.append((line_number, answer_value))
            continue
        clean_lines.append((line_number, line))

    return correct_letters, clean_lines


def _strip_inline_answer_key(line):
    lowered = _fold_quiz_text(line)
    positions = [lowered.find(prefix) for prefix in _ANSWER_PREFIXES if lowered.find(prefix) >= 0]
    if not positions:
        return line, set()

    marker_index = min(positions)
    answer_key = lowered[marker_index:].split(':', 1)[1] if ':' in lowered[marker_index:] else ''
    correct_letters = {letter for letter in ('a', 'b', 'c', 'd', 'e', 'f') if letter in answer_key}
    return line[:marker_index].strip(), correct_letters


def _split_inline_choices(line):
    line, inline_correct_letters = _strip_inline_answer_key(line)
    matches = list(_CHOICE_MARKER_RE.finditer(line))
    if len(matches) < 2:
        return None, [], inline_correct_letters

    question_text = line[:matches[0].start()].strip()
    answers = []

    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(line)
        answer_text = line[match.start():end].strip()
        if answer_text:
            answers.append(answer_text)

    return question_text, answers, inline_correct_letters


def _is_question_line(line):
    value = line.strip().lower()
    if not value:
        return False
    prefixes = ('câu ', 'cau ', 'question ', 'q.')
    if value.startswith(prefixes):
        return True
    if len(value) >= 3 and value[0].isdigit():
        marker_index = 1
        while marker_index < len(value) and value[marker_index].isdigit():
            marker_index += 1
        return marker_index < len(value) and value[marker_index] in '. ):'
    return False


def _is_section_heading(line):
    value = _fold_quiz_text(line).strip()
    while value and not value[0].isalnum():
        value = value[1:]
    return value.startswith('phan ') or value.startswith('section ')


def _is_choice_line(line):
    value = line.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    return bool(
        _CHOICE_LINE_RE.match(value)
        or _NUMERIC_CHOICE_LINE_RE.match(value)
        or _BULLET_CHOICE_LINE_RE.match(value)
        or _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value))
    )


def _choice_letter(answer, index):
    value = answer.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    named_match = _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value))
    if named_match:
        return named_match.group(1).lower()
    match = _CHOICE_LINE_RE.match(value)
    if match:
        for char in match.group(0):
            if char.lower() in 'abcdef':
                return char.lower()
    return chr(ord('a') + index)


def _clean_choice_prefix(answer):
    value = answer.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    if value.endswith('*'):
        value = value[:-1].strip()
    if _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value)):
        for separator in (':', '.', ')', '-', '/', chr(8211), chr(8212)):
            if separator in value:
                return value.split(separator, 1)[1].strip()
    value = _NUMERIC_CHOICE_PREFIX_RE.sub('', value, count=1).strip()
    value = _BULLET_CHOICE_PREFIX_RE.sub('', value, count=1).strip()
    return _CHOICE_LINE_RE.sub('', value, count=1).strip()


def _parse_question_block(first_line_number, question_lines, answer_lines):
    question_text = _normalize_quiz_text('\n'.join(question_lines))
    correct_letters, answer_lines = _extract_correct_letters(answer_lines)

    if not question_text:
        return None, f'Dòng {first_line_number}: câu hỏi không hợp lệ.'

    if not answer_lines:
        return None, f'Dòng {first_line_number}: câu hỏi chưa có đáp án.'

    choices = []
    correct_count = 0

    for index, (line_number, answer) in enumerate(answer_lines):
        raw_answer = answer.strip()
        is_correct = (
            raw_answer.startswith('*')
            or raw_answer.endswith('*')
            or _choice_letter(raw_answer, index) in correct_letters
        )
        choice_text = _normalize_quiz_text(_clean_choice_prefix(raw_answer))

        if not choice_text:
            return None, f'Dòng {line_number}: đáp án không được để trống.'

        if is_correct:
            correct_count += 1
        choices.append((choice_text, is_correct))

    if correct_count == 0:
        first_choice, _ = choices[0]
        choices[0] = (first_choice, True)

    return (question_text, choices), None



def _parse_answers_for_question(question_text, answers):
    """Parse answers from the manual add/edit forms with import rules."""
    answer_lines = [
        (line_number, line.strip())
        for line_number, line in enumerate(answers.splitlines(), start=1)
        if line.strip()
    ]
    parsed, error = _parse_question_block(1, [question_text], answer_lines)
    if error:
        return [], error
    return parsed[1], None


def _parse_uploaded_quiz_file(uploaded_file):
    text, errors = _extract_uploaded_quiz_text(uploaded_file)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    questions = []
    question_lines = []
    answer_lines = []
    first_line_number = None

    def flush_current():
        
        nonlocal question_lines, answer_lines, first_line_number
        if not question_lines and not answer_lines:
            return
        parsed, error = _parse_question_block(first_line_number or 1, question_lines, answer_lines)
        if parsed:
            questions.append(parsed)
        if error:
            errors.append(error)
        question_lines = []
        answer_lines = []
        first_line_number = None

    for line_number, raw_line in enumerate(text.split('\n'), start=1):
        line = raw_line.strip()
        if not line:
            if question_lines or answer_lines:
                flush_current()
            continue
        if line.startswith("'"):
            flush_current()
            continue

        lowered = _fold_quiz_text(line)
        is_answer_marker = lowered.startswith(_ANSWER_PREFIXES)
        inline_question, inline_answers, inline_correct_letters = _split_inline_choices(line)

        if _is_section_heading(line):
            if question_lines and answer_lines:
                flush_current()
            elif question_lines:
                question_lines = []
                first_line_number = None
            continue

        # A numbered line after choices starts the next question; check this
        # before numeric-choice recognition (e.g. "1. ...").
        if question_lines and answer_lines and _is_question_line(line):
            flush_current()

        if inline_answers and inline_question:
            if question_lines or answer_lines:
                flush_current()
            first_line_number = line_number
            question_lines.append(inline_question)
            answer_lines.extend((line_number, answer) for answer in inline_answers)
            if inline_correct_letters:
                answer_lines.append((line_number, "Đáp án: " + ",".join(sorted(inline_correct_letters))))
            continue
            
        if question_lines and (_is_choice_line(line) or is_answer_marker or line.startswith('*')):
            answer_lines.append((line_number, line))
            continue

        if _is_question_line(line) and (question_lines or answer_lines):
            flush_current()

        # Files exported as alternating question/answer rows often do not
        # label the answer with A., B. or a bullet. Treat the next line as
        # the answer when the question itself has an explicit number/label.
        if question_lines and not answer_lines and _is_question_line(question_lines[0]):
            answer_lines.append((line_number, line))
            continue

        if not question_lines:
            first_line_number = line_number
            question_lines.append(line)
            continue

        if _is_choice_line(line) or is_answer_marker or line.startswith('*'):
            answer_lines.append((line_number, line))
            continue

        if answer_lines:
            prev_line_number, prev_text = answer_lines[-1]
            answer_lines[-1] = (prev_line_number, prev_text + '\n' + line)
        else:
            question_lines.append(line)
    
    flush_current()

    if not questions and not errors:
        errors.append('File chưa có câu hỏi hoặc đáp án có thể nhận diện được.')

    return questions, errors

def _can_access_quiz(user, quiz):
    def _can_access_quiz(user, quiz):
        if user.is_admin:
            return True
        if quiz.created_by == user or not quiz.classroom_id:
            return True

    return ClassroomMember.objects.filter(
        classroom=quiz.classroom,
        student=user,
        status='approved'
    ).exists()

@login_required
def create_quiz(request):

    if request.method == 'POST':

        title = request.POST.get('title')
        description = request.POST.get('description')
        max_attempts = request.POST.get('max_attempts') or None
        time_limit = request.POST.get('time_limit') or None

        quiz = Quiz.objects.create(
            title=title,
            description=description,
            created_by=request.user,
            max_attempts=max_attempts,
            time_limit=time_limit
        )

        mode = request.POST.get('mode') or 'manual'
        if mode == 'ai':
            quiz_file = request.FILES.get('quiz_file')
            if not quiz_file:
                quiz.delete()
                messages.error(request, "Vui lòng tải file Word (.docx), PDF hoặc file văn bản có câu hỏi trước khi tạo đề tự động.")
                return render(request, 'quiz_ai/create_quiz.html')

            parsed_questions, parse_errors = _parse_uploaded_quiz_file(quiz_file)
            if not parsed_questions:
                quiz.delete()
                messages.error(request, "File chưa có câu hỏi đúng cấu trúc. Bài kiểm tra chưa được tạo.")
                for error in parse_errors[:10]:
                    messages.warning(request, error)
                return render(request, 'quiz_ai/create_quiz.html')

            question_images = _extract_uploaded_quiz_images(quiz_file)
            with transaction.atomic():
                for index, (content, choices) in enumerate(parsed_questions):
                    question = Question.objects.create(
                        quiz=quiz,
                        content=content,
                        image=question_images[index] if index < len(question_images) else None,
                    )
                    for choice_text, is_correct in choices:
                        Choice.objects.create(
                            question=question,
                            content=choice_text,
                            is_correct=is_correct
                        )

            messages.success(request, f"Đã tạo {len(parsed_questions)} câu hỏi từ file tải lên.")
            if parse_errors:
                messages.warning(request, f"Có {len(parse_errors)} câu hỏi/dòng sai cấu trúc và đã được bỏ qua.")
                for error in parse_errors[:10]:
                    messages.warning(request, error)
            return redirect('quiz_ai:quiz_detail', quiz.id)
        manual_questions = request.POST.get('manual_questions', '').strip()
        if manual_questions:
            for lines in _split_manual_question_blocks(manual_questions):
                if len(lines) < 2:
                    continue
                parsed, error = _parse_question_block(
                    1,
                    [lines[0]],
                    list(enumerate(lines[1:], start=2)),
                )
                if error:
                    messages.warning(request, error)
                    continue
                content, choices = parsed
                question = Question.objects.create(quiz=quiz, content=content)
                for choice_text, is_correct in choices:
                    Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)
            return redirect('quiz_ai:quiz_detail', quiz.id)

        return redirect('quiz_ai:add_question', quiz.id)

    return render(request, 'quiz_ai/create_quiz.html')


@login_required
def add_question(request, quiz_id):
    quiz = get_object_or_404(Quiz.objects.prefetch_related('questions__choices'), id=quiz_id)

    if not request.user.is_admin and  quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền chỉnh sửa bài kiểm tra này.")

    if request.method == 'POST':
        content = _normalize_quiz_text(request.POST.get('content', ''))
        answers = request.POST.get('answers', '').splitlines()

        choices, error = _parse_answers_for_question(content, '\n'.join(answers))
        if error:
            messages.error(request, error)
            return redirect('quiz_ai:add_question', quiz.id)

        question = Question.objects.create(
            quiz=quiz,
            content=content,
            image=request.FILES.get('image') or None,
        )
        for choice_text, is_correct in choices:
            Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)

        messages.success(request, "Đã thêm câu hỏi mới.")
        return redirect('quiz_ai:add_question', quiz.id)

    questions = quiz.questions.all()
    template = get_template("quiz_ai/add_question.html")
    print("TEMPLATE:", template.origin.name)
    return render(request, 'quiz_ai/add_question.html', {
        'quiz': quiz,
        'questions': questions
    })


@login_required
def quiz_list(request):
    keyword = request.GET.get('q', '').strip()
    quizzes = Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions')

    if keyword:
        quizzes = quizzes.filter(
            Q(title__icontains=keyword) |
            Q(description__icontains=keyword) |
            Q(classroom__name__icontains=keyword)
        )

    return render(request, 'quiz_ai/quiz_list.html', {
        'quizzes': quizzes.order_by('-created_at'),
        'keyword': keyword,
    })


@login_required
def quiz_detail(request, quiz_id):
    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and  _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền xem đề thi này.")

    questions = list(quiz.questions.all())
    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1

    attempt_count = QuizResult.objects.filter(user=request.user, quiz=quiz).count()
    blocked = bool(quiz.max_attempts and attempt_count >= quiz.max_attempts and quiz.created_by != request.user)

    return render(request, 'quiz_ai/quiz_detail.html', {
        'quiz': quiz,
        'questions': questions,
        'blocked': blocked,
        'attempt_count': attempt_count,
    })


@login_required
def take_quiz(request, quiz_id):
    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and  _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền làm bài kiểm tra này.")

    attempt_count = QuizResult.objects.filter(user=request.user, quiz=quiz).count()
    if quiz.max_attempts and attempt_count >= quiz.max_attempts and quiz.created_by != request.user:
        return render(request, 'quiz_ai/attempt_limit.html', {'quiz': quiz})

    question_order = request.GET.get('shuffle_questions')
    answer_order = request.GET.get('shuffle_answers')
    show_order_setup = question_order is None or answer_order is None
    shuffle_questions = question_order == '1'
    shuffle_answers = answer_order == '1'

    questions = list(quiz.questions.all())
    if shuffle_questions:
        random.shuffle(questions)

    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1
        question.display_choices = list(question.choices.all())
        if shuffle_answers:
            random.shuffle(question.display_choices)

    if not show_order_setup:
        request.session[f'quiz_start_{quiz.id}'] = timezone.now().isoformat()

    return render(request, 'quiz_ai/take_quiz.html', {
        'quiz': quiz,
        'questions': questions,
        'attempt_count': attempt_count,
        'show_order_setup': show_order_setup,
        'shuffle_questions': shuffle_questions,
        'shuffle_answers': shuffle_answers,
    })


@login_required
def submit_quiz(request, quiz_id):
    if request.method != 'POST':
        return redirect('quiz_ai:take_quiz', quiz_id=quiz_id)

    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền nộp bài kiểm tra này.")

    total_questions = quiz.questions.count()
    correct_answers = 0

    for question in quiz.questions.all():
        correct_ids = {
            str(choice.id)
            for choice in question.choices.all()
            if choice.is_correct
        }
        selected_ids = set(request.POST.getlist(f'question_{question.id}'))

        if correct_ids and selected_ids == correct_ids:
            correct_answers += 1

    score = round((correct_answers / total_questions) * 10) if total_questions else 0

    start_time_str = request.session.pop(f'quiz_start_{quiz.id}', None)
    total_seconds = 0
    if start_time_str:
        try:
            start_time = datetime.fromisoformat(start_time_str)
            total_seconds = max(0, int((timezone.now() - start_time).total_seconds()))
        except ValueError:
            total_seconds = 0

    minutes = total_seconds // 60
    seconds = total_seconds % 60

    QuizResult.objects.create(
        user=request.user,
        quiz=quiz,
        score=score
    )

    return render(request, 'quiz_ai/result.html', {
        'quiz': quiz,
        'score': score,
        'correct_answers': correct_answers,
        'total_questions': total_questions,
        'minutes': minutes,
        'seconds': seconds,
    })


@login_required
def edit_question(request, question_id):
    question = get_object_or_404(
        Question.objects.select_related('quiz').prefetch_related('choices'),
        id=question_id
    )

    if not request.user.is_admin and question.quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền chỉnh sửa câu hỏi này.")

    if request.method == 'POST':
        question.quiz.title = request.POST.get('title', '').strip()
        question.quiz.description = request.POST.get('description', '').strip()
        question.quiz.save(update_fields=['title', 'description'])

        question.content = _normalize_quiz_text(request.POST.get('content', ''))
        image = request.FILES.get('image')
        if request.POST.get('remove_image') and question.image:
            question.image.delete(save=False)
            question.image = None
        elif image:
            question.image = image
        question.save()

        choices, error = _parse_answers_for_question(question.content, request.POST.get('answers', ''))
        if error:
            messages.error(request, error)
            return redirect('quiz_ai:edit_question', question.id)

        question.choices.all().delete()
        for choice_text, is_correct in choices:
            Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)

        messages.success(request, "Đã lưu thay đổi câu hỏi.")
        return redirect('quiz_ai:add_question', question.quiz.id)

    return render(request, 'quiz_ai/edit_question.html', {'question': question})


@login_required
def delete_question(request, question_id):
    question = get_object_or_404(Question.objects.select_related('quiz'), id=question_id)
    quiz = question.quiz

    if not request.user.is_admin and quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa câu hỏi này.")

    if request.method == 'POST':
        question.delete()
        messages.success(request, "Đã xóa câu hỏi.")

    return redirect('quiz_ai:add_question', quiz.id)


@login_required
def delete_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    if not request.user.is_admin and quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa đề thi này.")

    if request.method == 'POST':
        quiz.delete()
        messages.success(request, "Đã xóa đề thi.")

    return redirect('quiz_ai:quiz_list')


@login_required
def flashcards(request, quiz_id):
    quiz = get_object_or_404(Quiz.objects.prefetch_related('questions__choices'), id=quiz_id)

    if not request.user.is_admin and _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền xem flashcards của đề này.")

    questions = list(quiz.questions.all())
    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1

    return render(request, 'quiz_ai/quiz_detail.html', {
        'quiz': quiz,
        'questions': questions,
        'blocked': False,
    })


@login_required
def quiz_history(request):
    results = QuizResult.objects.filter(
        Q(user=request.user) |
        Q(quiz__created_by=request.user)
    ).select_related(
        'user',
        'quiz',
        'quiz__created_by'
    ).distinct().order_by('-created_at')

    return render(request, 'quiz_ai/history.html', {'results': results})


@login_required
def delete_result(request, result_id):
    result = get_object_or_404(QuizResult, id=result_id)

    if not request.user.is_admin and result.quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa lịch sử này")

    if request.method == "POST":
        result.delete()
        messages.success(request, "Đã xóa lịch sử làm bài")

    return redirect('quiz_ai:quiz_history')




