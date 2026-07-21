from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Quiz, Question, Choice, QuizResult, QuizReloadPenalty
from django.http import HttpResponseForbidden, HttpResponse
from django.utils import timezone
from datetime import datetime
from django.contrib import messages
from django.db.models import Q
from django.db import transaction
from django.urls import reverse
from classroom.models import Classroom, ClassroomMember
from django.template.loader import get_template
from django.core.files.uploadedfile import SimpleUploadedFile
from django.conf import settings
import os
import cloudinary.uploader
import re
import unicodedata
import random
from io import BytesIO
from PIL import Image




_CHOICE_MARKER_RE = re.compile(r'(?<!\w)[\s\-*]*(?:\(?[A-Fa-f]\)?)[\s]*[\.\):\-/\u2013\u2014]\s*')
_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*(?:\(?[A-Fa-f]\)?)[\s]*(?:[\.\):\-/\u2013\u2014]|\s+)')
_NAMED_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*(?:dap an|phuong an|option|choice)\s+([a-f])\s*[\.\):\-/\u2013\u2014]\s+')
_NUMERIC_CHOICE_LINE_RE = re.compile(r'^[\s\-*]*[1-9][0-9]*[\.)]\s*\S+$')
_NUMERIC_CHOICE_PREFIX_RE = re.compile(r'^[\s\-*]*[1-9][0-9]*[\.)]\s*')
_BULLET_CHOICE_LINE_RE = re.compile(r'^[\s]*(?:[-\u2022\u2013\u2014]|\*)\s+\S+$')
_BULLET_CHOICE_PREFIX_RE = re.compile(r'^[\s]*(?:[-\u2022\u2013\u2014]|\*)\s+')
_ANSWER_PREFIXES = (
    'dap an:', 'dap an dung:', 'answer:', 'answer key:', 'correct:',
)
_DOCX_IMAGE_MARKER_RE = re.compile(r'\[\[DOCX_IMAGE_\d+\]\]')
_DOCX_IMAGE_EXTENSIONS = {
    'image/jpeg': 'jpg',
    'image/jpg': 'jpg',
    'image/png': 'png',
    'image/gif': 'gif',
    'image/webp': 'webp',
    'image/bmp': 'bmp',
    'image/x-wmf': 'wmf',
    'image/wmf': 'wmf',
    'image/x-emf': 'emf',
    'image/emf': 'emf',
}
_DOCX_CONVERTIBLE_IMAGE_TYPES = {'image/x-wmf', 'image/wmf', 'image/x-emf', 'image/emf'}


def _fold_quiz_text(value):
    normalized = unicodedata.normalize('NFD', value.lower().strip())
    without_marks = ''.join(ch for ch in normalized if unicodedata.category(ch) != 'Mn')
    return without_marks.replace(chr(273), 'd')


def _decode_uploaded_text(uploaded_file):
    raw_content = uploaded_file.read()
    for encoding in ('utf-8-sig', 'utf-8', 'cp1258', 'cp1252'):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_content.decode('utf-8', errors='replace')


def _normalize_quiz_text(value):
    return value.strip().replace('<br />', '\n').replace('<br/>', '\n').replace('<br>', '\n')



def _decode_uploaded_bytes(raw_content):
    for encoding in ('utf-8-sig', 'utf-8', 'cp1258', 'cp1252'):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_content.decode('utf-8', errors='replace')



def _extract_uploaded_quiz_images(uploaded_file):
    """Return embedded Word images in document order for automatic quizzes."""
    filename = (getattr(uploaded_file, 'name', '') or '').lower()
    if not filename.endswith('.docx'):
        return []

    try:
        from docx import Document
        raw_content = uploaded_file.read()
        uploaded_file.seek(0)
        document = Document(BytesIO(raw_content))
    except Exception:
        return []

    images = []
    for index, shape in enumerate(document.inline_shapes, start=1):
        try:
            relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = document.part.related_parts[relationship_id]
            images.append(_docx_image_file_from_part(image_part, index))
        except Exception:
            continue
    return images


def _store_word_question_image(image_file, quiz_id, index):
    """Upload a Word-extracted image explicitly and return its Cloudinary ID."""
    if not image_file:
        return None
    if not settings.USE_CLOUDINARY_MEDIA:
        return image_file

    image_file.seek(0)
    result = cloudinary.uploader.upload(
        image_file,
        resource_type="image",
        folder=f"teddy/quiz_question_images/quiz_{quiz_id}",
        public_id=f"question_{index + 1}",
        overwrite=True,
    )
    return result["public_id"]


def _combine_word_images(image_files, name):
    image_files = [image_file for image_file in image_files if image_file]
    if not image_files:
        return None
    if len(image_files) == 1:
        try:
            image_files[0].seek(0)
        except (AttributeError, OSError):
            pass
        return image_files[0]

    images = []
    for image_file in image_files:
        try:
            image_file.seek(0)
            with Image.open(image_file) as image:
                loaded = image.convert('RGBA')
                loaded.load()
                images.append(loaded)
        except Exception:
            continue

    if not images:
        return image_files[0]

    padding = 12
    width = max(image.width for image in images)
    height = sum(image.height for image in images) + padding * (len(images) - 1)
    canvas = Image.new('RGB', (width, height), 'white')
    y = 0
    for image in images:
        x = (width - image.width) // 2
        background = Image.new('RGB', image.size, 'white')
        if image.mode == 'RGBA':
            background.paste(image, mask=image.getchannel('A'))
        else:
            background.paste(image)
        canvas.paste(background, (x, y))
        y += image.height + padding

    output = BytesIO()
    canvas.save(output, format='PNG')
    return SimpleUploadedFile(
        name=f'{name}.png',
        content=output.getvalue(),
        content_type='image/png',
    )


def _word_images_for_markers(value, image_map):
    return [
        image_map[marker]
        for marker in _DOCX_IMAGE_MARKER_RE.findall(value or '')
        if marker in image_map
    ]


def _strip_docx_image_markers(value):
    return _DOCX_IMAGE_MARKER_RE.sub('', value or '').strip()


def _is_docx_image_only_line(value):
    return bool(_DOCX_IMAGE_MARKER_RE.search(value or '')) and not _strip_docx_image_markers(value)



_MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
_MATH_TAG = '{' + _MATH_NS + '}'


def _xml_attr(node, local_name):
    for key, value in node.attrib.items():
        if key.endswith('}' + local_name) or key == local_name:
            return value
    return ''


def _first_child_by_local_name(node, local_name):
    for child in node:
        if child.tag.endswith('}' + local_name):
            return child
    return None


def _math_group_text(node):
    if node is None:
        return ''
    return _docx_math_text(node).strip()


def _docx_math_text(node):
    tag = node.tag
    if tag.endswith('}t'):
        return node.text or ''

    if tag.endswith('}f'):
        numerator = _math_group_text(_first_child_by_local_name(node, 'num'))
        denominator = _math_group_text(_first_child_by_local_name(node, 'den'))
        if numerator and denominator:
            return f'({numerator})/({denominator})'

    if tag.endswith('}sSup'):
        base = _math_group_text(_first_child_by_local_name(node, 'e'))
        sup = _math_group_text(_first_child_by_local_name(node, 'sup'))
        if base and sup:
            return f'{base}^{sup}'

    if tag.endswith('}sSub'):
        base = _math_group_text(_first_child_by_local_name(node, 'e'))
        sub = _math_group_text(_first_child_by_local_name(node, 'sub'))
        if base and sub:
            return f'{base}_{sub}'

    if tag.endswith('}sSubSup'):
        base = _math_group_text(_first_child_by_local_name(node, 'e'))
        sub = _math_group_text(_first_child_by_local_name(node, 'sub'))
        sup = _math_group_text(_first_child_by_local_name(node, 'sup'))
        if base:
            suffix = (f'_{sub}' if sub else '') + (f'^{sup}' if sup else '')
            return base + suffix

    if tag.endswith('}rad'):
        degree = _math_group_text(_first_child_by_local_name(node, 'deg'))
        expression = _math_group_text(_first_child_by_local_name(node, 'e'))
        if expression:
            return f'{degree}√({expression})' if degree else f'√({expression})'

    if tag.endswith('}d'):
        props = _first_child_by_local_name(node, 'dPr')
        begin = end = ''
        if props is not None:
            begin_node = _first_child_by_local_name(props, 'begChr')
            end_node = _first_child_by_local_name(props, 'endChr')
            begin = _xml_attr(begin_node, 'val') if begin_node is not None else ''
            end = _xml_attr(end_node, 'val') if end_node is not None else ''
        expression = _math_group_text(_first_child_by_local_name(node, 'e'))
        return f'{begin or "("}{expression}{end or ")"}' if expression else ''

    if tag.endswith('}nary'):
        symbol = ''
        props = _first_child_by_local_name(node, 'naryPr')
        if props is not None:
            char_node = _first_child_by_local_name(props, 'chr')
            symbol = _xml_attr(char_node, 'val') if char_node is not None else ''
        sub = _math_group_text(_first_child_by_local_name(node, 'sub'))
        sup = _math_group_text(_first_child_by_local_name(node, 'sup'))
        expression = _math_group_text(_first_child_by_local_name(node, 'e'))
        prefix = symbol or '∑'
        if sub:
            prefix += f'_{sub}'
        if sup:
            prefix += f'^{sup}'
        return (prefix + ' ' + expression).strip()

    parts = []
    for child in node:
        value = _docx_math_text(child)
        if value:
            parts.append(value)
    return ''.join(parts)


def _docx_cell_text(cell, document, image_map):
    blocks = [_docx_paragraph_with_images(paragraph, document, image_map) for paragraph in cell.paragraphs]
    return '\n'.join(block for block in blocks if block).strip()


def _docx_image_file_from_part(image_part, image_index):
    content_type = getattr(image_part, 'content_type', '') or 'image/png'
    blob = image_part.blob
    extension = _DOCX_IMAGE_EXTENSIONS.get(content_type, content_type.rsplit('/', 1)[-1] or 'png')

    if content_type in _DOCX_CONVERTIBLE_IMAGE_TYPES:
        try:
            with Image.open(BytesIO(blob)) as image:
                output = BytesIO()
                image.save(output, format='PNG')
            blob = output.getvalue()
            content_type = 'image/png'
            extension = 'png'
        except Exception:
            pass

    return SimpleUploadedFile(
        name=f'word_image_{image_index}.{extension}',
        content=blob,
        content_type=content_type,
    )


def _docx_related_image_marker(relationship_id, document, image_map):
    if not relationship_id or relationship_id not in document.part.related_parts:
        return ''

    image_part = document.part.related_parts[relationship_id]
    content_type = getattr(image_part, 'content_type', '')
    if not content_type.startswith('image/'):
        return ''

    marker = f'[[DOCX_IMAGE_{len(image_map)}]]'
    image_map[marker] = _docx_image_file_from_part(image_part, len(image_map) + 1)
    return marker


def _docx_image_markers_from_node(node, document, image_map):
    markers = []
    seen_relationships = set()
    for image_node in node.iter():
        for key, relationship_id in image_node.attrib.items():
            local_name = key.rsplit('}', 1)[-1]
            if local_name not in {'embed', 'id', 'link'} or relationship_id in seen_relationships:
                continue
            marker = _docx_related_image_marker(relationship_id, document, image_map)
            if marker:
                markers.append(marker)
                seen_relationships.add(relationship_id)
    return markers


def _docx_paragraph_with_images(paragraph, document, image_map):
    """Keep Word text, Office Math text, and embedded equation/image previews in order."""
    parts = []
    for run in paragraph._p:
        if run.tag.startswith(_MATH_TAG):
            parts.append(_docx_math_text(run))
            continue
        if not run.tag.endswith('}r'):
            continue
        for node in run:
            if node.tag.startswith(_MATH_TAG):
                parts.append(_docx_math_text(node))
            elif node.tag.endswith('}t'):
                parts.append(node.text or '')
            elif node.tag.endswith('}tab'):
                parts.append('\t')
            elif node.tag.endswith(('}object', '}drawing', '}pict')):
                parts.extend(_docx_image_markers_from_node(node, document, image_map))
    return ''.join(parts).strip()


def _extract_docx_answer_key(document):
    answer_key = {}
    for table in document.tables:
        table_answer_key = {}
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for index in range(0, len(rows) - 1, 2):
            for number, answer in zip(rows[index], rows[index + 1]):
                if number.isdigit() and answer.upper() in {'A', 'B', 'C', 'D', 'E', 'F'}:
                    table_answer_key[int(number)] = answer.upper()
        for row in rows:
            for cell in row:
                folded = _fold_quiz_text(cell)
                for match in re.finditer(r'(?:cau\s*)?(\d{1,3})\s*[\.\):\-]?\s*([a-f])\b', folded):
                    table_answer_key[int(match.group(1))] = match.group(2).upper()
        if len(table_answer_key) >= 2:
            answer_key.update(table_answer_key)
    return answer_key


def _line_has_numbered_answer_pairs(line):
    folded = _fold_quiz_text(line)
    return bool(re.search(r'(?:cau\s*)?\d{1,3}\s*[\.\):\-]?\s*[a-f]\b', folded))


def _is_answer_key_heading(line):
    folded = _fold_quiz_text(line).strip(' :.-').strip()
    return (
        folded in {'dap an', 'dap an dung', 'answer', 'answer key', 'key'}
        or folded.startswith('bang dap an')
        or folded.startswith('answer key')
    )


def _extract_answer_key_from_text(text):
    answer_key = {}
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    in_key = False

    for index, line in enumerate(lines):
        folded = _fold_quiz_text(line)
        starts_key = _is_answer_key_heading(line) or ('dap an' in folded and _line_has_numbered_answer_pairs(line))
        if starts_key:
            in_key = True
        if not in_key:
            continue

        for match in re.finditer(r'(?:cau\s*)?(\d{1,3})\s*[\.\):\-]?\s*([a-f])\b', folded):
            answer_key.setdefault(int(match.group(1)), set()).add(match.group(2).lower())

        if index + 1 < len(lines):
            numbers = re.findall(r'\b\d{1,3}\b', line)
            letters = re.findall(r'\b[A-Fa-f]\b', lines[index + 1])
            if len(numbers) >= 2 and len(numbers) == len(letters):
                for number, letter in zip(numbers, letters):
                    answer_key.setdefault(int(number), set()).add(letter.lower())

    return answer_key


def _apply_answer_key_to_questions(questions, answer_key):
    resolved_questions = []
    for number, (section, content, choices) in enumerate(questions, start=1):
        if choices and not any(is_correct for _, is_correct in choices):
            correct_letters = answer_key.get(number, set())
            if correct_letters:
                choices = [
                    (choice_text, chr(ord('a') + index) in correct_letters)
                    for index, (choice_text, _) in enumerate(choices)
                ]
        if choices and not any(is_correct for _, is_correct in choices):
            first_choice, _ = choices[0]
            choices[0] = (first_choice, True)
        resolved_questions.append((section, content, choices))
    return resolved_questions

def _extract_uploaded_quiz_text(uploaded_file):
    raw_content = uploaded_file.read()
    try:
        uploaded_file.seek(0)
    except (AttributeError, OSError):
        pass

    filename = (getattr(uploaded_file, 'name', '') or '').lower()
    extension = filename.rsplit('.', 1)[-1] if '.' in filename else ''

    if extension in {'txt', 'text', 'csv'}:
        return _decode_uploaded_bytes(raw_content), []

    if extension == 'docx':
        try:
            from docx import Document
        except ImportError:
            return '', ['Máy chủ chưa có thư viện đọc Word. Vui lòng cài python-docx.']
        try:
            document = Document(BytesIO(raw_content))
        except Exception:
            return '', ['Không thể đọc tệp Word. Vui lòng tải tệp .docx hợp lệ.']

        image_map = {}
        blocks = [_docx_paragraph_with_images(paragraph, document, image_map) for paragraph in document.paragraphs]
        blocks = [block for block in blocks if block]
        uploaded_file._docx_image_map = image_map
        uploaded_file._docx_answer_key = _extract_docx_answer_key(document)
        for table in document.tables:
            for row in table.rows:
                cells = [_docx_cell_text(cell, document, image_map) for cell in row.cells]
                if len(cells) < 2 or not cells[0] or not cells[1]:
                    continue
                header = _fold_quiz_text(f'{cells[0]} {cells[1]}')
                if 'cau hoi' in header and ('dap an' in header or 'lua chon' in header):
                    continue
                blocks.append(f'{cells[0]}\n{cells[1]}')
        # Each Word paragraph is a consecutive quiz line. Separating every
        # paragraph with a blank line flushes the parser before its choices.
        return '\n'.join(blocks), []

    if extension == 'pdf':
        try:
            import pdfplumber
        except ImportError:
            return '', ['Máy chủ chưa có thư viện đọc PDF. Vui lòng cài pdfplumber.']
        try:
            with pdfplumber.open(BytesIO(raw_content)) as pdf:
                plain_text = []
                for page in pdf.pages:
                    simple_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ''
                    layout_text = page.extract_text(layout=True) or ''
                    page_text = simple_text if simple_text.strip() else layout_text
                    if page_text.strip():
                        plain_text.append(page_text)
        except Exception:
            return '', ['Không thể đọc tệp PDF. Vui lòng tải PDF hợp lệ hoặc chuyển sang Word (.docx).']

        text = '\n\n'.join(plain_text).strip()
        if not text:
            return '', ['PDF này có vẻ là file scan/ảnh nên không có văn bản để quét. Vui lòng dùng PDF có thể chọn/sao chép chữ hoặc chuyển sang Word (.docx).']
        return text, []

    return '', ['Chỉ hỗ trợ tệp .docx, .pdf, .txt, .text hoặc .csv.']


def _split_manual_question_blocks(text):
    normalized = text.replace('\r\n', '\n').replace('\r', '\n')
    blocks = []
    current = []
    current_section = ''

    def flush_current():
        nonlocal current
        if current:
            blocks.append((current_section, current))
            current = []

    for line in normalized.split('\n'):
        stripped = line.strip()
        if not stripped:
            flush_current()
            continue
        section = _section_title_from_marker(stripped)
        if section:
            flush_current()
            current_section = section
            continue
        if stripped.startswith("'"):
            flush_current()
            continue
        current.append(stripped)

    flush_current()
    return blocks


def _extract_correct_letters(lines):
    correct_letters = set()
    clean_lines = []

    for line_number, line in lines:
        lowered = _fold_quiz_text(line)
        if lowered.startswith(_ANSWER_PREFIXES):
            answer_value = line.split(':', 1)[1].strip() if ':' in line else ''
            compact_key = re.sub(r'[\s,;/&]+', '', _fold_quiz_text(answer_value))
            compact_key = compact_key.replace('va', '').replace('and', '')
            if compact_key and all(letter in 'abcdef' for letter in compact_key):
                correct_letters.update(compact_key)
                continue
            
            clean_lines.append((line_number, answer_value))
            continue
        clean_lines.append((line_number, line))

    return correct_letters, clean_lines


def _strip_inline_answer_key(line):
    lowered = _fold_quiz_text(line)
    positions = [lowered.find(prefix) for prefix in _ANSWER_PREFIXES if lowered.find(prefix) >= 0]
    if not positions:
        return line, set()

    marker_index = min(positions)
    answer_key = lowered[marker_index:].split(':', 1)[1] if ':' in lowered[marker_index:] else ''
    correct_letters = {letter for letter in ('a', 'b', 'c', 'd', 'e', 'f') if letter in answer_key}
    return line[:marker_index].strip(), correct_letters


def _split_inline_choices(line):
    line, inline_correct_letters = _strip_inline_answer_key(line)
    matches = list(_CHOICE_MARKER_RE.finditer(line))
    if len(matches) < 2:
        return None, [], inline_correct_letters

    question_text = line[:matches[0].start()].strip()
    answers = []

    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(line)
        answer_text = line[match.start():end].strip()
        if answer_text:
            answers.append(answer_text)

    return question_text, answers, inline_correct_letters


def _is_question_line(line):
    value = line.strip().lower()
    if not value:
        return False
    prefixes = ('câu ', 'cau ', 'question ', 'q.')
    if value.startswith(prefixes):
        return True
    if len(value) >= 3 and value[0].isdigit():
        marker_index = 1
        while marker_index < len(value) and value[marker_index].isdigit():
            marker_index += 1
        return marker_index < len(value) and value[marker_index] in '. ):'
    return False


def _is_section_heading(line):
    value = _fold_quiz_text(line).strip()
    while value and not value[0].isalnum():
        value = value[1:]
    return value.startswith('phan ') or value.startswith('section ')



def _section_title_from_marker(line):
    value = line.strip()
    if value.startswith("'"):
        value = value[1:].strip()
    return value if _is_section_heading(value) else ''


def _mark_section_headings(questions):
    last_section = None
    for question in questions:
        section = (getattr(question, 'section', '') or '').strip()
        question.show_section_heading = bool(section and section != last_section)
        if section:
            last_section = section
    return questions

def _is_choice_line(line):
    value = line.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    return bool(
        _CHOICE_LINE_RE.match(value)
        or _NUMERIC_CHOICE_LINE_RE.match(value)
        or _BULLET_CHOICE_LINE_RE.match(value)
        or _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value))
    )


def _choice_letter(answer, index):
    value = answer.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    named_match = _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value))
    if named_match:
        return named_match.group(1).lower()
    match = _CHOICE_LINE_RE.match(value)
    if match:
        for char in match.group(0):
            if char.lower() in 'abcdef':
                return char.lower()
    return chr(ord('a') + index)


def _clean_choice_prefix(answer):
    value = answer.strip()
    if value.startswith('*'):
        value = value[1:].strip()
    if value.endswith('*'):
        value = value[:-1].strip()
    if _NAMED_CHOICE_LINE_RE.match(_fold_quiz_text(value)):
        for separator in (':', '.', ')', '-', '/', chr(8211), chr(8212)):
            if separator in value:
                return value.split(separator, 1)[1].strip()
    value = _NUMERIC_CHOICE_PREFIX_RE.sub('', value, count=1).strip()
    value = _BULLET_CHOICE_PREFIX_RE.sub('', value, count=1).strip()
    return _CHOICE_LINE_RE.sub('', value, count=1).strip()


def _parse_question_block(first_line_number, question_lines, answer_lines, default_first_correct=True):
    while answer_lines:
        first_answer = answer_lines[0][1].strip()
        first_folded = _fold_quiz_text(first_answer)
        first_is_explicit_answer = (
            _is_choice_line(first_answer)
            or first_folded.startswith(_ANSWER_PREFIXES)
            or first_answer.startswith('*')
        )
        has_later_explicit_answer = any(
            _is_choice_line(answer)
            or _fold_quiz_text(answer).startswith(_ANSWER_PREFIXES)
            or answer.strip().startswith('*')
            for _, answer in answer_lines[1:]
        )
        if first_is_explicit_answer or not has_later_explicit_answer:
            break
        question_lines.append(answer_lines.pop(0)[1])

    question_text = _normalize_quiz_text('\n'.join(question_lines))
    correct_letters, answer_lines = _extract_correct_letters(answer_lines)
    expanded_answers = []
    for line_number, answer in answer_lines:
        markers = list(_CHOICE_MARKER_RE.finditer(answer))
        if len(markers) >= 2:
            for index, marker in enumerate(markers):
                end = markers[index + 1].start() if index + 1 < len(markers) else len(answer)
                expanded_answers.append((line_number, answer[marker.start():end].strip()))
        else:
            expanded_answers.append((line_number, answer))
    answer_lines = expanded_answers

    if not question_text:
        return None, f'Dòng {first_line_number}: câu hỏi không hợp lệ.'

    if not answer_lines:
        return None, f'Dòng {first_line_number}: câu hỏi chưa có đáp án.'

    choices = []
    correct_count = 0

    for index, (line_number, answer) in enumerate(answer_lines):
        raw_answer = answer.strip()
        is_correct = (
            raw_answer.startswith('*')
            or raw_answer.endswith('*')
            or _choice_letter(raw_answer, index) in correct_letters
        )
        choice_text = _normalize_quiz_text(_clean_choice_prefix(raw_answer))

        if not choice_text:
            return None, f'Dòng {line_number}: đáp án không được để trống.'

        if is_correct:
            correct_count += 1
        choices.append((choice_text, is_correct))

    if correct_count == 0 and default_first_correct:
        first_choice, _ = choices[0]
        choices[0] = (first_choice, True)

    return (question_text, choices), None



def _parse_answers_for_question(question_text, answers):
    """Parse answers from the manual add/edit forms with import rules."""
    answer_lines = [
        (line_number, line.strip())
        for line_number, line in enumerate(answers.splitlines(), start=1)
        if line.strip()
    ]
    parsed, error = _parse_question_block(1, [question_text], answer_lines)
    if error:
        return [], error
    return parsed[1], None


def _parse_uploaded_quiz_file(uploaded_file):
    text, errors = _extract_uploaded_quiz_text(uploaded_file)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    answer_key = _extract_answer_key_from_text(text)
    for number, letter in getattr(uploaded_file, '_docx_answer_key', {}).items():
        try:
            answer_key.setdefault(int(number), set()).add(str(letter).lower())
        except (TypeError, ValueError):
            continue

    questions = []
    question_lines = []
    answer_lines = []
    first_line_number = None
    current_section = ''

    def flush_current():
        nonlocal question_lines, answer_lines, first_line_number
        if not question_lines and not answer_lines:
            return
        parsed, error = _parse_question_block(
            first_line_number or 1,
            question_lines,
            answer_lines,
            default_first_correct=False,
        )
        if parsed:
            content, choices = parsed
            questions.append((current_section, content, choices))
        if error:
            errors.append(error)
        question_lines = []
        answer_lines = []
        first_line_number = None

    for line_number, raw_line in enumerate(text.split('\n'), start=1):
        line = raw_line.strip()
        if not line:
            if answer_lines:
                flush_current()
            continue

        section = _section_title_from_marker(line)
        if section:
            flush_current()
            current_section = section
            continue
        if line.startswith("'"):
            flush_current()
            continue

        lowered = _fold_quiz_text(line)
        if _is_answer_key_heading(line) or ('dap an' in lowered and _line_has_numbered_answer_pairs(line)):
            flush_current()
            break

        is_answer_marker = lowered.startswith(_ANSWER_PREFIXES)
        inline_question, inline_answers, inline_correct_letters = _split_inline_choices(line)

        if _is_section_heading(line):
            flush_current()
            current_section = line
            continue

        if question_lines and answer_lines and _is_question_line(line):
            flush_current()

        if inline_answers and inline_question:
            if question_lines or answer_lines:
                flush_current()
            first_line_number = line_number
            question_lines.append(inline_question)
            answer_lines.extend((line_number, answer) for answer in inline_answers)
            if inline_correct_letters:
                answer_lines.append((line_number, "Đáp án: " + ",".join(sorted(inline_correct_letters))))
            continue

        if question_lines and (_is_choice_line(line) or is_answer_marker or line.startswith('*')):
            answer_lines.append((line_number, line))
            continue

        if _is_question_line(line) and (question_lines or answer_lines):
            flush_current()

        if _is_docx_image_only_line(line):
            if answer_lines:
                prev_line_number, prev_text = answer_lines[-1]
                answer_lines[-1] = (prev_line_number, prev_text + '\n' + line)
            else:
                if not question_lines:
                    first_line_number = line_number
                question_lines.append(line)
            continue

        if question_lines and not answer_lines and _is_question_line(question_lines[0]):
            answer_lines.append((line_number, line))
            continue

        if not question_lines:
            first_line_number = line_number
            question_lines.append(line)
            continue

        if _is_choice_line(line) or is_answer_marker or line.startswith('*'):
            answer_lines.append((line_number, line))
            continue

        if answer_lines:
            prev_line_number, prev_text = answer_lines[-1]
            answer_lines[-1] = (prev_line_number, prev_text + '\n' + line)
        else:
            question_lines.append(line)

    flush_current()
    questions = _apply_answer_key_to_questions(questions, answer_key)

    if not questions and not errors:
        errors.append('File chưa có câu hỏi hoặc đáp án có thể nhận diện được.')

    return questions, errors



def _build_take_question_sections(questions):
    has_sections = any((getattr(question, 'section', '') or '').strip() for question in questions)
    if not has_sections:
        for number, question in enumerate(questions, start=1):
            question.display_number = number
        return [{'title': '', 'questions': questions, 'index': 1}], False

    sections = []
    section_map = {}
    for question in questions:
        title = (getattr(question, 'section', '') or '').strip() or 'Câu hỏi chung'
        if title not in section_map:
            section = {'title': title, 'questions': [], 'index': len(sections) + 1}
            sections.append(section)
            section_map[title] = section
        section_map[title]['questions'].append(question)

    display_number = 1
    for section in sections:
        for section_number, question in enumerate(section['questions'], start=1):
            question.display_number = display_number
            question.section_number = section_number
            display_number += 1
    return sections, True

def _quiz_attempt_count(user, quiz):
    return (
        QuizResult.objects.filter(user=user, quiz=quiz).count()
        + QuizReloadPenalty.objects.filter(user=user, quiz=quiz).count()
    )

def _can_access_quiz(user, quiz):
    if user.is_admin:
        return True
    if quiz.created_by == user or not quiz.classroom_id:
        return True
    return ClassroomMember.objects.filter(
        classroom=quiz.classroom,
        student=user,
        status='approved'
    ).exists()
@login_required
def create_quiz(request):
    classroom_id = request.POST.get('classroom_id') or request.GET.get('classroom_id')
    classroom = None
    if classroom_id:
        classroom = get_object_or_404(Classroom, id=classroom_id)
        if not request.user.is_admin and classroom.teacher != request.user:
            return HttpResponseForbidden("Bạn không có quyền tạo bài kiểm tra trong lớp này.")

    if request.method == 'POST':

        title = request.POST.get('title')
        description = request.POST.get('description')
        max_attempts = request.POST.get('max_attempts') or None
        time_limit = request.POST.get('time_limit') or None

        quiz = Quiz.objects.create(
            title=title,
            description=description,
            created_by=request.user,
            classroom=classroom,
            max_attempts=max_attempts,
            time_limit=time_limit
        )

        mode = request.POST.get('mode') or 'manual'
        if mode == 'ai':
            quiz_file = request.FILES.get('quiz_file')
            if not quiz_file:
                quiz.delete()
                messages.error(request, "Vui lòng tải file Word (.docx), PDF hoặc file văn bản có câu hỏi trước khi tạo đề tự động.")
                return render(request, 'quiz/create_quiz.html', {'selected_classroom': classroom})

            parsed_questions, parse_errors = _parse_uploaded_quiz_file(quiz_file)
            if not parsed_questions:
                quiz.delete()
                messages.error(request, "File chưa có câu hỏi đúng cấu trúc. Bài kiểm tra chưa được tạo.")
                for error in parse_errors[:10]:
                    messages.warning(request, error)
                return render(request, 'quiz/create_quiz.html', {'selected_classroom': classroom})

            image_map = getattr(quiz_file, '_docx_image_map', {})
            with transaction.atomic():
                for index, (section, content, choices) in enumerate(parsed_questions):
                    image = _combine_word_images(
                        _word_images_for_markers(content, image_map),
                        f'question_{index + 1}',
                    )
                    if image:
                        image = _store_word_question_image(image, quiz.id, index)
                    clean_content = _strip_docx_image_markers(content)
                    question = Question.objects.create(
                        quiz=quiz,
                        content=clean_content,
                        section=section,
                        image=image,
                    )
                    for choice_index, (choice_text, is_correct) in enumerate(choices):
                        choice_image = _combine_word_images(
                            _word_images_for_markers(choice_text, image_map),
                            f'question_{index + 1}_choice_{choice_index + 1}',
                        )
                        if choice_image:
                            choice_image = _store_word_question_image(choice_image, quiz.id, index * 10 + choice_index)
                        Choice.objects.create(
                            question=question,
                            content=_strip_docx_image_markers(choice_text),
                            is_correct=is_correct,
                            image=choice_image,
                        )

            messages.success(request, f"Đã tạo {len(parsed_questions)} câu hỏi từ file tải lên.")
            if parse_errors:
                messages.warning(request, f"Có {len(parse_errors)} câu hỏi/dòng sai cấu trúc và đã được bỏ qua.")
                for error in parse_errors[:10]:
                    messages.warning(request, error)
            return redirect('quiz:quiz_detail', quiz.id)
        manual_questions = request.POST.get('manual_questions', '').strip()
        if manual_questions:
            manual_blocks = _split_manual_question_blocks(manual_questions)
            for section, lines in manual_blocks:
                if len(lines) < 2:
                    continue
                parsed, error = _parse_question_block(
                    1,
                    [lines[0]],
                    list(enumerate(lines[1:], start=2)),
                )
                if error:
                    messages.warning(request, error)
                    continue
                content, choices = parsed
                question = Question.objects.create(quiz=quiz, content=content, section=section)
                for choice_text, is_correct in choices:
                    Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)
            return redirect('quiz:quiz_detail', quiz.id)

        return redirect('quiz:add_question', quiz.id)

    return render(request, 'quiz/create_quiz.html', {'selected_classroom': classroom})


@login_required
def add_question(request, quiz_id):
    quiz = get_object_or_404(Quiz.objects.prefetch_related('questions__choices'), id=quiz_id)

    if not request.user.is_admin and  quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền chỉnh sửa bài kiểm tra này.")

    if request.method == 'POST':
        content = _normalize_quiz_text(request.POST.get('content', ''))
        answers = request.POST.get('answers', '').splitlines()

        choices, error = _parse_answers_for_question(content, '\n'.join(answers))
        if error:
            messages.error(request, error)
            return redirect('quiz:add_question', quiz.id)

        question = Question.objects.create(
            quiz=quiz,
            content=content,
            section=request.POST.get('section', '').strip(),
            image=request.FILES.get('image') or None,
        )
        for choice_text, is_correct in choices:
            Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)

        messages.success(request, "Đã thêm câu hỏi mới.")
        return redirect('quiz:add_question', quiz.id)

    questions = _mark_section_headings(list(quiz.questions.all()))
    template = get_template("quiz/add_question.html")
    print("TEMPLATE:", template.origin.name)
    return render(request, 'quiz/add_question.html', {
        'quiz': quiz,
        'questions': questions
    })


@login_required
def quiz_list(request):
    keyword = request.GET.get('q', '').strip()
    quizzes = Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions')

    if keyword:
        quizzes = quizzes.filter(
            Q(title__icontains=keyword) |
            Q(description__icontains=keyword) |
            Q(classroom__name__icontains=keyword)
        )

    return render(request, 'quiz/quiz_list.html', {
        'quizzes': quizzes.order_by('-created_at'),
        'keyword': keyword,
    })


@login_required
def quiz_detail(request, quiz_id):
    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and not _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền xem đề thi này.")

    questions = _mark_section_headings(list(quiz.questions.all()))
    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1

    attempt_count = _quiz_attempt_count(request.user, quiz)
    blocked = bool(quiz.max_attempts and attempt_count >= quiz.max_attempts and quiz.created_by != request.user)

    return render(request, 'quiz/quiz_detail.html', {
        'quiz': quiz,
        'questions': questions,
        'blocked': blocked,
        'attempt_count': attempt_count,
    })


@login_required
def take_quiz(request, quiz_id):
    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by', 'classroom').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and not _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền làm bài kiểm tra này.")

    attempt_count = _quiz_attempt_count(request.user, quiz)
    reload_key = f'quiz_active_{quiz.id}'
    start_key = f'quiz_start_{quiz.id}'

    question_order = request.GET.get('shuffle_questions')
    answer_order = request.GET.get('shuffle_answers')
    show_order_setup = question_order is None or answer_order is None
    shuffle_questions = question_order == '1'
    shuffle_answers = answer_order == '1'
    start_requested = not show_order_setup

    if quiz.max_attempts and attempt_count >= quiz.max_attempts and quiz.created_by != request.user:
        request.session.pop(reload_key, None)
        request.session.pop(start_key, None)
        return render(request, 'quiz/attempt_limit.html', {'quiz': quiz})

    if start_requested and request.session.get(reload_key):
        request.session.pop(reload_key, None)
        request.session.pop(start_key, None)
        return redirect('quiz:take_quiz', quiz_id=quiz.id)

    if start_requested:
        if quiz.created_by != request.user:
            QuizReloadPenalty.objects.create(user=request.user, quiz=quiz)
            attempt_count += 1
        request.session[reload_key] = True
        request.session[start_key] = timezone.now().isoformat()

    questions = list(quiz.questions.all())
    has_sections = any((getattr(question, 'section', '') or '').strip() for question in questions)
    if shuffle_questions:
        if has_sections:
            grouped_questions = []
            section_groups = {}
            for question in questions:
                section_title = (question.section or '').strip() or 'Câu hỏi chung'
                section_groups.setdefault(section_title, []).append(question)
            for group in section_groups.values():
                random.shuffle(group)
                grouped_questions.extend(group)
            questions = grouped_questions
        else:
            random.shuffle(questions)
    question_sections, has_question_sections = _build_take_question_sections(questions)

    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1
        question.display_choices = list(question.choices.all())
        if shuffle_answers:
            random.shuffle(question.display_choices)

    return render(request, 'quiz/take_quiz.html', {
        'quiz': quiz,
        'questions': questions,
                'question_sections': question_sections,
        'has_question_sections': has_question_sections,
        'attempt_count': attempt_count,
        'show_order_setup': show_order_setup,
        'shuffle_questions': shuffle_questions,
        'shuffle_answers': shuffle_answers,
    })

@login_required
def submit_quiz(request, quiz_id):
    if request.method != 'POST':
        return redirect('quiz:take_quiz', quiz_id=quiz_id)

    quiz = get_object_or_404(
        Quiz.objects.select_related('created_by').prefetch_related('questions__choices'),
        id=quiz_id
    )

    if not request.user.is_admin and not _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền nộp bài kiểm tra này.")

    total_questions = quiz.questions.count()
    correct_answers = 0

    for question in quiz.questions.all():
        correct_ids = {
            str(choice.id)
            for choice in question.choices.all()
            if choice.is_correct
        }
        selected_ids = set(request.POST.getlist(f'question_{question.id}'))

        if correct_ids and selected_ids == correct_ids:
            correct_answers += 1

    score = round((correct_answers / total_questions) * 10) if total_questions else 0

    active_key = f'quiz_active_{quiz.id}'
    if request.session.get(active_key):
        current_attempt = QuizReloadPenalty.objects.filter(
            user=request.user,
            quiz=quiz,
        ).order_by('-created_at').first()
        if current_attempt:
            current_attempt.delete()

    request.session.pop(active_key, None)
    start_time_str = request.session.pop(f'quiz_start_{quiz.id}', None)
    total_seconds = 0
    if start_time_str:
        try:
            start_time = datetime.fromisoformat(start_time_str)
            total_seconds = max(0, int((timezone.now() - start_time).total_seconds()))
        except ValueError:
            total_seconds = 0

    minutes = total_seconds // 60
    seconds = total_seconds % 60

    QuizResult.objects.create(
        user=request.user,
        quiz=quiz,
        score=score
    )

    return render(request, 'quiz/result.html', {
        'quiz': quiz,
        'score': score,
        'correct_answers': correct_answers,
        'total_questions': total_questions,
        'minutes': minutes,
        'seconds': seconds,
    })


@login_required
def edit_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    if not request.user.is_admin and quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền chỉnh sửa đề thi này.")

    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        if not title:
            messages.error(request, "Tên đề thi không được để trống.")
            return redirect('quiz:edit_quiz', quiz.id)

        try:
            time_limit = int(request.POST.get('time_limit') or 0) or None
            max_attempts = int(request.POST.get('max_attempts') or 0) or None
            if (time_limit and time_limit < 1) or (max_attempts and max_attempts < 1):
                raise ValueError
        except ValueError:
            messages.error(request, "Thời gian và số lượt làm bài phải là số nguyên dương.")
            return redirect('quiz:edit_quiz', quiz.id)

        quiz.title = title
        quiz.description = request.POST.get('description', '').strip()
        quiz.time_limit = time_limit
        quiz.max_attempts = max_attempts
        quiz.save(update_fields=['title', 'description', 'time_limit', 'max_attempts'])
        messages.success(request, "Đã cập nhật thông tin đề thi.")
        return redirect('quiz:edit_quiz', quiz.id)

    return render(request, 'quiz/edit_quiz.html', {'quiz': quiz})

@login_required
def edit_question(request, question_id):
    question = get_object_or_404(
        Question.objects.select_related('quiz').prefetch_related('choices'),
        id=question_id
    )

    if not request.user.is_admin and question.quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền chỉnh sửa câu hỏi này.")

    if request.method == 'POST':
        question.content = _normalize_quiz_text(request.POST.get('content', ''))
        question.section = request.POST.get('section', '').strip()
        image = request.FILES.get('image')
        if request.POST.get('remove_image') and question.image:
            question.image.delete(save=False)
            question.image = None
        elif image:
            question.image = image
        question.save()

        choices, error = _parse_answers_for_question(question.content, request.POST.get('answers', ''))
        if error:
            messages.error(request, error)
            return redirect('quiz:edit_question', question.id)

        question.choices.all().delete()
        for choice_text, is_correct in choices:
            Choice.objects.create(question=question, content=choice_text, is_correct=is_correct)

        messages.success(request, "Đã lưu thay đổi câu hỏi.")
        return redirect('quiz:add_question', question.quiz.id)

    return render(request, 'quiz/edit_question.html', {'question': question})


@login_required
def delete_question(request, question_id):
    question = get_object_or_404(Question.objects.select_related('quiz'), id=question_id)
    quiz = question.quiz

    if not request.user.is_admin and quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa câu hỏi này.")

    if request.method == 'POST':
        question.delete()
        messages.success(request, "Đã xóa câu hỏi.")

    return redirect('quiz:add_question', quiz.id)


@login_required
def delete_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    if not request.user.is_admin and quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa đề thi này.")

    if request.method == 'POST':
        quiz.delete()
        messages.success(request, "Đã xóa đề thi.")

    return redirect('quiz:quiz_list')



@login_required
def download_quiz(request, quiz_id, format):
    quiz = get_object_or_404(Quiz.objects.prefetch_related('questions__choices'), id=quiz_id)

    filename = re.sub(r'[^\w-]+', '_', quiz.title, flags=re.UNICODE).strip('_') or 'de_kiem_tra'
    questions = list(quiz.questions.all())
    if format == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(quiz.title, 0)
        if quiz.description:
            document.add_paragraph(quiz.description)
        last_section = None
        for number, question in enumerate(questions, start=1):
            section = (question.section or '').strip()
            if section and section != last_section:
                document.add_heading(section, level=1)
                last_section = section
            document.add_heading(f'Câu {number}: {question.content}', level=2)
            for letter, choice in zip('ABCDEFGHIJKLMNOPQRSTUVWXYZ', question.choices.all()):
                marker = ' (Đáp án đúng)' if choice.is_correct else ''
                document.add_paragraph(f'{letter}. {choice.content}{marker}')
        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    if format == 'pdf':
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            from xml.sax.saxutils import escape
        except ImportError:
            return HttpResponse('PDF chưa sẵn sàng. Vui lòng tải bản Word.', status=503)

        font_regular = 'Helvetica'
        font_bold = 'Helvetica-Bold'
        font_candidates = [
            (r'C:\Windows\Fonts\DejaVuSans.ttf', r'C:\Windows\Fonts\DejaVuSans-Bold.ttf'),
            (r'C:\Windows\Fonts\arial.ttf', r'C:\Windows\Fonts\arialbd.ttf'),
        ]
        for regular_path, bold_path in font_candidates:
            if os.path.exists(regular_path) and os.path.exists(bold_path):
                pdfmetrics.registerFont(TTFont('TeddyUnicode', regular_path))
                pdfmetrics.registerFont(TTFont('TeddyUnicode-Bold', bold_path))
                font_regular = 'TeddyUnicode'
                font_bold = 'TeddyUnicode-Bold'
                break

        def paragraph_text(value):
            return escape(str(value or '')).replace('\n', '<br/>')

        buffer = BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
            title=quiz.title,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'QuizTitle',
            parent=styles['Title'],
            fontName=font_bold,
            fontSize=18,
            leading=24,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=8,
        )
        description_style = ParagraphStyle(
            'QuizDescription',
            parent=styles['BodyText'],
            fontName=font_regular,
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#475569'),
            spaceAfter=12,
        )
        section_style = ParagraphStyle(
            'QuizSection',
            parent=styles['Heading1'],
            fontName=font_bold,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#075985'),
            spaceBefore=14,
            spaceAfter=6,
        )
        question_style = ParagraphStyle(
            'QuizQuestion',
            parent=styles['BodyText'],
            fontName=font_bold,
            fontSize=11,
            leading=16,
            textColor=colors.HexColor('#111827'),
            spaceBefore=10,
            spaceAfter=4,
            keepWithNext=True,
        )
        choice_style = ParagraphStyle(
            'QuizChoice',
            parent=styles['BodyText'],
            fontName=font_regular,
            fontSize=10,
            leading=15,
            leftIndent=8 * mm,
            firstLineIndent=-5 * mm,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=3,
        )
        correct_choice_style = ParagraphStyle(
            'QuizCorrectChoice',
            parent=choice_style,
            fontName=font_bold,
            textColor=colors.HexColor('#166534'),
        )

        story = [Paragraph(paragraph_text(quiz.title), title_style)]
        if quiz.description:
            story.append(Paragraph(paragraph_text(quiz.description), description_style))

        last_section = None
        for number, question in enumerate(questions, start=1):
            section = (question.section or '').strip()
            if section and section != last_section:
                story.append(Paragraph(paragraph_text(section), section_style))
                last_section = section
            story.append(Paragraph(f'Câu {number}: {paragraph_text(question.content)}', question_style))
            for letter, choice in zip('ABCDEFGHIJKLMNOPQRSTUVWXYZ', question.choices.all()):
                suffix = ' (Đáp án đúng)' if choice.is_correct else ''
                style = correct_choice_style if choice.is_correct else choice_style
                story.append(Paragraph(f'{letter}. {paragraph_text(choice.content)}{suffix}', style))
            story.append(Spacer(1, 4))

        document.build(story)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
    return HttpResponse(status=404)

@login_required
def flashcards(request, quiz_id):
    quiz = get_object_or_404(Quiz.objects.prefetch_related('questions__choices'), id=quiz_id)

    if not request.user.is_admin and not _can_access_quiz(request.user, quiz):
        return HttpResponseForbidden("Bạn không có quyền xem flashcards của đề này.")

    questions = _mark_section_headings(list(quiz.questions.all()))
    for question in questions:
        question.allows_multiple = sum(1 for choice in question.choices.all() if choice.is_correct) > 1

    return render(request, 'quiz/flashcards.html', {
        'quiz': quiz,
        'questions': questions,
    })


@login_required
def quiz_history(request):
    submitted_results = QuizResult.objects.filter(
        Q(user=request.user) |
        Q(quiz__created_by=request.user)
    ).select_related(
        'user',
        'quiz',
        'quiz__created_by'
    ).distinct()

    reload_attempts = QuizReloadPenalty.objects.filter(
        Q(user=request.user) |
        Q(quiz__created_by=request.user)
    ).select_related(
        'user',
        'quiz',
        'quiz__created_by'
    ).distinct()

    records = []
    for result in submitted_results:
        records.append({
            'user': result.user,
            'quiz': result.quiz,
            'score': result.score,
            'created_at': result.created_at,
            'kind': 'submitted',
            'status_label': 'Đã nộp bài',
            'delete_url': reverse('quiz:delete_result', args=[result.id]),
        })

    for attempt in reload_attempts:
        records.append({
            'user': attempt.user,
            'quiz': attempt.quiz,
            'score': None,
            'created_at': attempt.created_at,
            'kind': 'reload',
            'status_label': 'Tải lại / chưa nộp',
            'delete_url': reverse('quiz:delete_reload_attempt', args=[attempt.id]),
        })

    records.sort(key=lambda record: record['created_at'], reverse=True)
    return render(request, 'quiz/history.html', {'results': records})


@login_required
def delete_result(request, result_id):
    result = get_object_or_404(QuizResult, id=result_id)

    if not request.user.is_admin and result.quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa lịch sử này")

    if request.method == "POST":
        result.delete()
        messages.success(request, "Đã xóa lịch sử làm bài")

    return redirect('quiz:quiz_history')


@login_required
def delete_reload_attempt(request, penalty_id):
    attempt = get_object_or_404(QuizReloadPenalty.objects.select_related('quiz'), id=penalty_id)

    if not request.user.is_admin and attempt.quiz.created_by != request.user:
        return HttpResponseForbidden("Bạn không có quyền xóa lịch sử này")

    if request.method == "POST":
        attempt.delete()
        messages.success(request, "Đã xóa lịch sử tải lại bài làm")

    return redirect('quiz:quiz_history')
